import streamlit as st
import google.generativeai as genai
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from streamlit_gsheets import GSheetsConnection
import datetime

# --- 1. НАСТРОЙКИ СТРАНИЦЫ (Должны быть в самом верху) ---
st.set_page_config(page_title="Methodist PRO", layout="wide", page_icon="📚")

# --- ДАННЫЕ АВТОРА ---
AUTHOR_NAME = "Адильбаева Айнура Дуйшембековна"
INSTAGRAM_HANDLE = "uchitel_tdk"
INSTAGRAM_URL = f"https://instagram.com/{INSTAGRAM_HANDLE}"
WHATSAPP_URL = "https://wa.me/77776513022"
PHONE_NUMBER = "+7 (777) 651-30-22"

# --- 2. СЛОВАРЬ ПЕРЕВОДОВ ---
TRANS = {
    "login_title": {"RU": "Вход в систему Методист PRO", "KZ": "Methodist PRO жүйесіне кіру"},
    "login_prompt": {"RU": "Введите ваш номер телефона для доступа.", "KZ": "Кіру үшін телефон нөміріңізді енгізіңіз."},
    "phone_label": {"RU": "Номер телефона:", "KZ": "Телефон нөмірі:"},
    "login_btn": {"RU": "Войти", "KZ": "Кіру"},
    "access_denied": {"RU": "Доступ закрыт. Номер не найден.", "KZ": "Кіруге тыйым салынды. Нөмір табылмады."},
    "status_active": {"RU": "✅ Подписка PRO активна", "KZ": "✅ PRO жазылым белсенді"},
    "teacher_fio": {"RU": "ФИО Учителя:", "KZ": "Мұғалімнің А.Т.Ә.:"},
    "date_label": {"RU": "Дата урока:", "KZ": "Сабақ күні:"},
    "subject_label": {"RU": "Предмет:", "KZ": "Пән:"},
    "grade_label": {"RU": "Класс:", "KZ": "Сынып:"},
    "topic_label": {"RU": "Тема урока:", "KZ": "Сабақтың тақырыбы:"},
    "score_label": {"RU": "Макс. балл:", "KZ": "Макс. ұпай:"},
    "goals_label": {"RU": "Цели обучения (ЦО):", "KZ": "Оқу мақсаттары (ОМ):"},
    "ksp_goals": {"RU": "Цели урока:", "KZ": "Сабақтың мақсаты:"},
    "mat_type": {"RU": "Тип материала:", "KZ": "Материал түрі:"},
    "type_work": {"RU": "Рабочий лист", "KZ": "Жұмыс парағы"},
    "type_sor": {"RU": "БЖБ (СОР) / ТЖБ (СОЧ)", "KZ": "БЖБ (СОР) / ТЖБ (СОЧ)"},
    "tab_class": {"RU": "📝 ЗАДАНИЯ (СОР/СОЧ)", "KZ": "📝 ТАПСЫРМАЛАР (БЖБ/ТЖБ)"},
    "tab_inc": {"RU": "👤 ИНКЛЮЗИЯ (Отдельно)", "KZ": "👤 ЕРЕКШЕ БІЛІМ (Жеке)"},
    "tab_ksp": {"RU": "📖 КСП (130 приказ)", "KZ": "📖 ҚМЖ (130-бұйрық)"},
    "inc_check": {"RU": "Есть ученик с ООП (Инклюзия)?", "KZ": "Ерекше білім беру қажеттілігі бар оқушы бар ма?"},
    "inc_diag": {"RU": "Диагноз/Особенности:", "KZ": "Диагноз/Ерекшеліктері:"},
    "func_lit": {"RU": "🧠 Функциональная грамотность (PISA)", "KZ": "🧠 Функционалдық сауаттылық (PISA)"},
    "btn_create": {"RU": "🚀 Создать материал", "KZ": "🚀 Материал жасау"},
    "download_btn": {"RU": "💾 СКАЧАТЬ WORD", "KZ": "💾 WORD ЖҮКТЕУ"},
    "preview": {"RU": "### Предпросмотр:", "KZ": "### Алдын ала қарау:"},
    "exit_btn": {"RU": "Выйти", "KZ": "Шығу"},
    "auth_title": {"RU": "Автор", "KZ": "Автор"}
}

# --- СПИСКИ ПРЕДМЕТОВ ---
SUBJECTS_RU = [
    "Русский язык (Я1 - родной)", "Русский язык (Я2 - второй)", 
    "Казахский язык (Т1 - родной)", "Казахский язык (Т2 - второй)",
    "Литературное чтение", "Обучение грамоте", "Букварь", "Ана тілі",
    "Математика", "Алгебра", "Геометрия", 
    "Естествознание", "Познание мира", 
    "Физика", "Химия", "Биология", "География", "Информатика",
    "История Казахстана", "Всемирная история", 
    "Английский язык", 
    "Музыка", "Художественный труд", "Изобразительное искусство (Рисование)", "Физическая культура"
]

SUBJECTS_KZ = [
    "Орыс тілі (Я1 - орыс сыныптары)", "Орыс тілі (Я2 - қазақ сыныптары)", 
    "Қазақ тілі (Т1 - қазақ сыныптары)", "Қазақ тілі (Т2 - орыс сыныптары)",
    "Әдебиеттік оқу", "Сауат ашу", "Әліппе", "Ана тілі",
    "Математика", "Алгебра", "Геометрия", 
    "Жаратылыстану", "Дүниетану", 
    "Физика", "Химия", "Биология", "География", "Информатика",
    "Қазақстан тарихы", "Дүниежүзі тарихы", 
    "Ағылшын тілі", 
    "Музыка", "Көркем еңбек", "Бейнелеу өнері", "Дене шынықтыру"
]

def get_text(key, lang_code):
    return TRANS.get(key, {}).get(lang_code, key)

# --- 3. АВТОРИЗАЦИЯ И ИИ (ИСПРАВЛЕНО) ---
def check_access(user_phone):
    try:
        conn = st.connection("gsheets", type=GSheetsConnection)
        # Добавляем обработку ошибок, если таблица недоступна
        df = conn.read(spreadsheet=st.secrets["gsheet_url"], ttl=0)
        clean_input = ''.join(filter(str.isdigit, str(user_phone)))
        # Предполагаем, что номера во 2-й колонке (индекс 1)
        allowed_phones = df.iloc[:, 1].astype(str).str.replace(r'\D', '', regex=True).tolist()
        return clean_input in allowed_phones
    except Exception as e: 
        st.error(f"Ошибка проверки доступа: {e}")
        return False

def configure_ai():
    try:
        # Получаем ключ из secrets.toml
        api_key = st.secrets.get("GOOGLE_API_KEY")
        if not api_key:
            st.error("API Key не найден в secrets!")
            return None
            
        genai.configure(api_key=api_key)
        
        # Используем самую стабильную версию модели на данный момент
        # gemini-1.5-flash обычно работает надежнее всего
        return genai.GenerativeModel('gemini-1.5-flash')
    except Exception as e:
        st.error(f"Ошибка подключения к AI: {e}")
        return None

# --- 4. ЛОГИКА ВХОДА ---
if 'lang' not in st.session_state: st.session_state['lang'] = 'RU'
if 'auth' not in st.session_state: st.session_state['auth'] = False

with st.sidebar:
    lang_select = st.selectbox("🌐 Тіл / Язык", ["Русский", "Қазақша"], index=0 if st.session_state['lang']=='RU' else 1)
    st.session_state['lang'] = "RU" if lang_select == "Русский" else "KZ"
    current_lang = st.session_state['lang']

if not st.session_state['auth']:
    st.title(get_text("login_title", current_lang))
    st.markdown(get_text("login_prompt", current_lang))
    phone_input = st.text_input(get_text("phone_label", current_lang))
    if st.button(get_text("login_btn", current_lang)):
        with st.spinner("Проверка..."):
            if check_access(phone_input):
                st.session_state['auth'] = True
                st.rerun()
            else: st.error(get_text("access_denied", current_lang))
    st.divider()
    st.caption(f"Dev: {AUTHOR_NAME}")
    st.stop()

# Инициализация модели после успешного входа
model = configure_ai()

# --- 5. БОКОВАЯ ПАНЕЛЬ ---
with st.sidebar:
    st.divider()
    st.success(get_text('status_active', current_lang))
    t_fio = st.text_input(get_text("teacher_fio", current_lang), value="Teacher")
    st.divider()
    st.markdown(f"### 👩‍💻 {get_text('auth_title', current_lang)}")
    st.info(f"**{AUTHOR_NAME}**")
    col1, col2 = st.columns(2)
    with col1: st.markdown(f"[![Inst](https://img.shields.io/badge/Inst-E4405F?logo=instagram&logoColor=white)]({INSTAGRAM_URL})")
    with col2: st.markdown(f"[![WA](https://img.shields.io/badge/WA-25D366?logo=whatsapp&logoColor=white)]({WHATSAPP_URL})")
    st.caption(f"📞 {PHONE_NUMBER}")
    if st.button(get_text("exit_btn", current_lang)):
        st.session_state['auth'] = False
        st.rerun()

# --- 6. ФУНКЦИЯ WORD ---
def clean_markdown(text):
    text = re.sub(r'[*_]{1,3}', '', text)
    text = re.sub(r'^#+\s*', '', text)
    return text.strip()

def create_docx(ai_text, title, subj, gr, teacher, lang_code, date_str, is_ksp=False, std_name=""):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Шапка
    labels = {
        "RU": {"student": "Ученик", "subj": "Предмет", "class": "Класс", "date": "Дата"},
        "KZ": {"student": "Оқушы", "subj": "Пән", "class": "Сынып", "date": "Күні"}
    }
    L = labels[lang_code]

    if not is_ksp:
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = f"{L['student']}: {std_name if std_name else '________________'}"
        table.cell(1, 0).text = f"{L['subj']}: {subj} | {L['class']}: {gr}"
        table.cell(0, 1).text = f"{L['date']}: {date_str}"
        doc.add_paragraph()

    h = doc.add_heading(title.upper(), 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0,0,0)
        run.font.size = Pt(14)
        run.bold = True
    
    lines = ai_text.split('\n')
    table_data = []
    
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|'):
            if '---' in stripped: continue
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            if cells: table_data.append(cells)
        else:
            if table_data:
                # Рисуем таблицу (с защитой от ошибок)
                cols_count = len(table_data[0])
                tbl = doc.add_table(rows=len(table_data), cols=cols_count)
                tbl.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    safe_cols = min(len(row), cols_count)
                    for j in range(safe_cols):
                        cell = tbl.cell(i, j)
                        cell.text = clean_markdown(row[j])
                        if i == 0: # Жирный заголовок
                            for p in cell.paragraphs:
                                for r in p.runs: r.font.bold = True
                table_data = []
                doc.add_paragraph()
            
            clean_line = clean_markdown(stripped)
            if clean_line:
                p = doc.add_paragraph(clean_line)
                # Жирный шрифт для ключевых слов
                keywords = ["задание", "тапсырма", "этап", "кезең", "критерии", "дескриптор", "ресурсы", "ответы", "жауаптар"]
                if any(clean_line.lower().startswith(x) for x in keywords):
                    if p.runs: p.runs[0].bold = True

    # Если таблица в конце
    if table_data:
        cols_count = len(table_data[0])
        tbl = doc.add_table(rows=len(table_data), cols=cols_count)
        tbl.style = 'Table Grid'
        for i, row in enumerate(table_data):
            safe_cols = min(len(row), cols_count)
            for j in range(safe_cols):
                tbl.cell(i, j).text = clean_markdown(row[j])

    doc.add_paragraph("\n" + "_"*30)
    doc.add_paragraph(f"{'Мұғалім' if lang_code=='KZ' else 'Учитель'}: {teacher}")
    doc.add_paragraph("Generated by Methodist PRO")
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- 7. ЦЕНТРАЛЬНАЯ ПАНЕЛЬ ---
st.title("🇰🇿 Methodist PRO")

# Глобальная дата
c_d1, c_d2 = st.columns([1, 4])
with c_d1:
    sel_date = st.date_input(get_text("date_label", current_lang), datetime.date.today())
    date_str = sel_date.strftime("%d.%m.%Y")

t1, t2, t3 = st.tabs([get_text("tab_class", current_lang), get_text("tab_inc", current_lang), get_text("tab_ksp", current_lang)])
subj_list = SUBJECTS_KZ if current_lang == "KZ" else SUBJECTS_RU

# === ВКЛАДКА 1: СОР/СОЧ/ЗАДАНИЯ ===
with t1:
    c1, c2, c3 = st.columns(3)
    with c1:
        m_subj = st.selectbox(get_text("subject_label", current_lang), subj_list, key="t1_s")
        m_grade = st.selectbox(get_text("grade_label", current_lang), [str(i) for i in range(1, 12)], key="t1_g")
    with c2:
        m_topic = st.text_input(get_text("topic_label", current_lang), key="t1_t")
        m_type = st.radio(get_text("mat_type", current_lang), [get_text("type_work", current_lang), get_text("type_sor", current_lang)], key="t1_type")
    with c3:
        m_score = st.number_input(get_text("score_label", current_lang), 1, 80, 10, key="t1_sc")
        # ГАЛОЧКА PISA
        use_pisa = st.checkbox(get_text("func_lit", current_lang), key="t1_pisa")
        
    m_goals = st.text_area(get_text("goals_label", current_lang), height=100, key="t1_gl")

    if st.button(get_text("btn_create", current_lang), type="primary", key="btn_t1"):
        if not m_goals.strip(): st.warning("Нет целей!")
        elif model is None: st.error("Ошибка: ИИ модель не настроена (проверьте API ключ).")
        else:
            lang_instr = "Пиши на КАЗАХСКОМ языке" if current_lang == "KZ" else "Пиши на РУССКОМ языке"
            pisa_instr = "Включи задания на функциональную грамотность (PISA)." if use_pisa else ""
            
            prompt = f"""
            Ты методист. {lang_instr}.
            Создай: {m_type}. Предмет: {m_subj}. Класс: {m_grade}. Тема: {m_topic}.
            Цели: {m_goals}. Макс балл: {m_score}.
            {pisa_instr}
            
            СТРУКТУРА:
            1. Задания разного уровня (A, B, C).
            2. ОБЯЗАТЕЛЬНО Таблица: "Критерии оценивания" и "Дескрипторы" (баллы).
            3. Ответы.
            """
            with st.spinner("Generating..."):
                try:
                    res = model.generate_content(prompt)
                    st.markdown(res.text)
                    doc = create_docx(res.text, m_topic, m_subj, m_grade, t_fio, current_lang, date_str, False)
                    st.download_button(get_text("download_btn", current_lang), doc, file_name=f"Task_{m_topic}.docx")
                except Exception as e: st.error(f"Error: {e}")

# === ВКЛАДКА 2: ИНКЛЮЗИЯ ===
with t2:
    st.info("Адаптация для ООП")
    ic1, ic2 = st.columns(2)
    with ic1:
        i_name = st.text_input("Имя ученика / Оқушының аты:", key="i_n")
        i_diag = st.text_input("Диагноз / Ерекшеліктері:", placeholder="Например: ЗПР, нарушение зрения", key="i_d")
    with ic2:
        i_topic = st.text_input("Тема (из первой вкладки):", value=m_topic, key="i_t")
        i_goals = st.text_area("Цели (упрощенные):", value=m_goals, height=100, key="i_g")

    if st.button("🧩 Адаптировать / Бейімдеу", type="primary", key="btn_t2"):
        if not i_goals: st.warning("Нет целей!")
        elif model is None: st.error("Ошибка: ИИ модель не настроена.")
        else:
            lang_instr = "Пиши на КАЗАХСКОМ" if current_lang == "KZ" else "Пиши на РУССКОМ"
            prompt = f"""
            Ты дефектолог. {lang_instr}.
            Адаптируй задания темы '{i_topic}' для ученика: {i_name}. Диагноз: {i_diag}.
            Цели: {i_goals}. Сделай задания проще, короче, понятнее.
            """
            with st.spinner("Adapting..."):
                try:
                    res = model.generate_content(prompt)
                    st.markdown(res.text)
                    doc = create_docx(res.text, f"Inclusion_{i_name}", m_subj, m_grade, t_fio, current_lang, date_str, False, i_name)
                    st.download_button(get_text("download_btn", current_lang), doc, file_name=f"Inc_{i_name}.docx")
                except Exception as e: st.error(f"Error: {e}")

# === ВКЛАДКА 3: КСП (130 ПРИКАЗ) ===
with t3:
    k1, k2 = st.columns(2)
    with k1:
        k_subj = st.selectbox(get_text("subject_label", current_lang), subj_list, key="k_s")
        k_grade = st.selectbox(get_text("grade_label", current_lang), [str(i) for i in range(1, 12)], key="k_g")
    with k2:
        k_topic = st.text_input(get_text("topic_label", current_lang), key="k_t")
        k_vals = st.text_input("Ценности / Құндылықтар:", value="Патриотизм, еңбекқорлық", key="k_v")

    k_om = st.text_area(get_text("goals_label", current_lang), placeholder="Код (например 5.1.2.1)...", key="k_om")
    k_sm = st.text_area(get_text("ksp_goals", current_lang), placeholder="Все учащиеся смогут...", key="k_sm")
    
    st.markdown("---")
    c_k1, c_k2 = st.columns(2)
    with c_k1:
        # ИНКЛЮЗИЯ В КСП
        use_inc = st.checkbox(get_text("inc_check", current_lang), key="k_inc_check")
        if use_inc:
            k_inc_desc = st.text_input(get_text("inc_diag", current_lang), placeholder="Пример: ЗПР", key="k_inc_inp")
    with c_k2:
        # PISA В КСП
        use_pisa_ksp = st.checkbox(get_text("func_lit", current_lang) + " (в КСП)", key="k_pisa_ksp")

    if st.button(get_text("btn_create", current_lang), type="primary", key="btn_ksp"):
        if not k_om.strip(): st.warning("Нет целей!")
        elif model is None: st.error("Ошибка: ИИ модель не настроена.")
        else:
            lang_instr = "Пиши на КАЗАХСКОМ" if current_lang == "KZ" else "Пиши на РУССКОМ"
            
            # Формируем структуру таблицы
            inc_col_header = ""
            inc_prompt = ""
            if use_inc:
                inc_col_header = "| Адаптация (ООП)"
                inc_prompt = f"В классе ученик с ООП ({k_inc_desc}). Добавь в таблицу столбец 'Адаптация' с упрощенными заданиями для него."
            
            pisa_prompt = "Включи активные методы и задания PISA." if use_pisa_ksp else ""

            prompt = f"""
            Ты методист (Казахстан, приказ 130). {lang_instr}.
            Составь КСП. Предмет: {k_subj}. Класс: {k_grade}. Тема: {k_topic}.
            ЦО: {k_om}. Цели урока: {k_sm}. Ценности: {k_vals}.
            {inc_prompt}
            {pisa_prompt}
            
            СТРУКТУРА ТАБЛИЦЫ (строго, используй Markdown таблицы):
            Этап урока | Действия педагога | Действия ученика {inc_col_header} | Оценивание | Ресурсы
            
            Этапы:
            1. Начало.
            2. Середина (Новая тема).
            3. Конец (Рефлексия).
            """
            
            with st.spinner("Generating Plan..."):
                try:
                    res = model.generate_content(prompt)
                    st.markdown(res.text)
                    doc = create_docx(res.text, f"КСП_{k_topic}", k_subj, k_grade, t_fio, current_lang, date_str, True)
                    st.download_button(get_text("download_btn", current_lang), doc, file_name=f"KSP_{k_topic}.docx")
                except Exception as e: st.error(f"Error: {e}")

st.markdown("---")
st.markdown(f"<center>{AUTHOR_NAME} © 2026</center>", unsafe_allow_html=True)
